import os
import urllib.request as request
from bs4 import BeautifulSoup
from docx import Document
from docx.oxml.ns import qn
import random

page_index = 0

def get_pages(url):
    item_text = ''
    response = request.urlopen(url)
    html = response.read()
    soup = BeautifulSoup(html)
    #print(soup)
    class_questions_col = soup.find('div', attrs={'class','questions_col'})
    list_li = class_questions_col.find_all('li')
    for item_li in list_li:
        table_content = item_li.find('table')
        if table_content == None:
            try:
                item_text += item_li.get_text() + '\n'
            except:
                continue
        else:
            for item_li_content in item_li.contents:
                if item_li_content == table_content:
                    for each_tag in item_li_content.contents[0].contents:
                        try:
                            item_text += each_tag.contents[0].string
                        except:
                            for td in each_tag.contents:
                                item_text += td.contents[0].string
                else:
                    try:
                        item_text += item_li_content.string
                    except:
                        continue

    item_text = item_text.replace('查看解析 添加到组卷', '').replace(' ', '').replace('\n\n\n\n','').replace('<br/>','')
    if item_text == '' or '暂时还没有资料' in item_text:
        print(page_index)
        print(item_text)
        return
    file_object = open('E:/数据/math_html/' + str(page_index) + '.txt', 'wb')
    file_object.write(item_text.encode('utf-8'))
    file_object.close()

def is_chinese(uchar):
    """判断一个unicode是否是汉字"""
    if uchar >= u'\u4e00' and uchar <= u'\u9fa5':
        return True
    else:
        return False

def is_number(uchar):
    """判断一个unicode是否是数字"""
    if uchar >= u'\u0030' and uchar <= u'\u0039':
        return True
    else:
        return False


def is_alphabet(uchar):
    """判断一个unicode是否是英文字母"""
    if (uchar >= u'\u0041' and uchar <= u'\u005a') or (uchar >= u'\u0061' and uchar <= u'\u007a'):
        return True
    else:
        return False

def txt2word(txt_path):
    html_dir = os.listdir(txt_path)

    document = Document()
    # document.styles['Normal'].font.name = u'Times New Roman'
    tag = 0
    for file_path in html_dir:
        tag += 1
        # if tag >= 100:
        #     break
        fopen = open(txt_path + file_path, 'rb')
        content_lines = fopen.read().decode('utf-8')
        for eachLine in content_lines.split('\n'):
            p = document.add_paragraph()
            if eachLine == '':
                continue
            if '暂时还没有资料' in eachLine:
                print(file_path)
                break
            for each_word in eachLine:
                if is_alphabet(each_word) or each_word == 'Φ' or each_word == 'ϕ' or each_word == 'φ':
                    run = p.add_run(each_word + '       ')
                    run.font.name = u'Times New Roman'
                    run.italic = True
                else:
                    run = p.add_run(each_word + '    ')
                    if each_word == '_':
                        run.font.bold = True
                        run.italic = True
                    run.font.name = u'宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    document.save('./data/word_data.docx')



if __name__ == "__main__":
    txt_path = './data/alltxt_input/'
    txt2word(txt_path)
    # for i in range(0,99):
    #     try:
    #         page_index = i
    #         get_pages("http://tiku.21cnjy.com/tiku.php?mod=quest&channel=3&xd=2&page=" + str(i+1))
    #     except:
    #         print(i)
    #         continue
